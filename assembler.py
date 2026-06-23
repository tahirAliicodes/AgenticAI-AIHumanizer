# assembler.py
# Stitches humanized chunks back into a proper document

from docx import Document
from docx.shared import Pt
import os

from db import save_job  # lets us write a row to the database


def assemble_txt(chunks, output_path):
    """Rebuilds a .txt file from humanized chunks."""
    lines = []

    for chunk in chunks:
        # Use humanized text if available, otherwise original
        text = chunk.get("humanized", chunk["text"])
        lines.append(text)

    # Join paragraphs with a blank line between them
    final_text = "\n\n".join(lines)

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(final_text)

    return output_path


def assemble_docx(chunks, output_path):
    """Rebuilds a .docx file from humanized chunks."""
    doc = Document()

    for chunk in chunks:
        text = chunk.get("humanized", chunk["text"])

        if chunk["type"] == "heading":
            # Re-apply the original heading style (Heading 1, Heading 2, etc.)
            level = chunk["style"].replace("Heading ", "")
            doc.add_heading(text, level=int(level))
        else:
            # Regular paragraph
            doc.add_paragraph(text)

    doc.save(output_path)
    return output_path


def assemble_document(chunks, original_filepath, output_dir, original_filename=None):
    """
    Main function — detects file type and assembles accordingly.

    original_filepath: where the uploaded file is saved on disk (may have a random unique name)
    original_filename: NEW — the real filename the user uploaded (e.g. "myessay.txt").
                        If not given, falls back to original_filepath's name (old behavior).
    """

    # NEW: prefer the real filename if we were given one
    if original_filename is None:
        original_filename = os.path.basename(original_filepath)

    name, ext = os.path.splitext(original_filename)
    output_filename = f"{name}_humanized{ext}"
    output_path = os.path.join(output_dir, output_filename)

    # Count total words across all humanized chunks
    word_count = sum(
        len(chunk.get("humanized", chunk["text"]).split())
        for chunk in chunks
    )

    if ext == ".txt":
        result = assemble_txt(chunks, output_path)
    elif ext == ".docx":
        result = assemble_docx(chunks, output_path)
    else:
        # Log the failed attempt before raising the error
        save_job(original_filename, word_count, status="failed")
        raise ValueError(f"Unsupported file type: {ext}")

    # Log this successful job to the database — using the REAL filename now
    save_job(original_filename, word_count, status="success")

    return result


# Quick test
if __name__ == "__main__":
    test_chunks = [
        {"type": "heading", "style": "Heading 1", "text": "Introduction", "humanized": "Introduction"},
        {"type": "paragraph", "text": "Old text.", "humanized": "This is the naturally rewritten version of the paragraph."},
        {"type": "paragraph", "text": "More old text.", "humanized": "And here is another humanized paragraph that flows nicely."},
    ]

    output = assemble_document(test_chunks, "test.docx", ".", original_filename="myessay.docx")
    print(f"Assembled document saved to: {output}")